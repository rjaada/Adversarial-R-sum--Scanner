"""
Extract plain text from PDF or DOCX. Returns text, parse_integrity score,
and a stripped ATS-preview string (what a naive ATS parser would see).
"""
from __future__ import annotations
import io
import re


def extract_resume_text(file_bytes: bytes, filename: str) -> dict:
    """
    Returns:
        text: full extracted text
        ats_preview: stripped version simulating naive ATS parse
        parse_integrity: float 0-1 (1 = no structural issues detected)
        warnings: list of parser warnings
    """
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        return _extract_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        return _extract_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {filename}")


def _detect_columns(words: list[dict]) -> bool:
    """
    Detect a multi-column layout from LINE-START x positions.

    The previous heuristic bucketed the x0 of EVERY word, which any normal
    full-width text line exceeds — it flagged virtually every PDF as
    multi-column (false 0.20 parse-integrity penalty + a false parse_warning
    issue on clean single-column résumés; confirmed by fixture measurement).

    Columns manifest as two (or more) dominant left edges where many lines
    begin, separated by a wide horizontal gap. So: group words into visual
    lines, take each line's leftmost x0, cluster those starts into 25pt
    buckets, and report columns only when at least two buckets each anchor
    3+ lines and sit >= 100pt apart (bullet indents are ~10-15pt and never
    trip this; a true second column sits far right).
    """
    if not words:
        return False
    line_starts: dict[int, float] = {}  # top-bucket (~4pt tolerance) -> min x0
    for w in words:
        key = round(w["top"] / 4)
        x = w["x0"]
        if key not in line_starts or x < line_starts[key]:
            line_starts[key] = x
    if len(line_starts) < 6:
        return False  # too few lines to judge — be conservative
    buckets: dict[int, int] = {}
    for x in line_starts.values():
        b = round(x / 25)
        buckets[b] = buckets.get(b, 0) + 1
    dominant = [b for b, count in buckets.items() if count >= 3]
    if len(dominant) < 2:
        return False
    return (max(dominant) - min(dominant)) * 25 >= 100


def _extract_pdf(file_bytes: bytes) -> dict:
    import pdfplumber

    warnings: list[str] = []
    pages_text: list[str] = []
    column_detected = False
    table_detected = False
    image_count = 0
    page_count = 0

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            text = page.extract_text(x_tolerance=2, y_tolerance=2) or ""
            pages_text.append(text)

            # Detect multi-column layouts (common ATS failure)
            if _detect_columns(page.extract_words()):
                column_detected = True

            # Detect tables (often lost in ATS)
            if page.extract_tables():
                table_detected = True

            image_count += len(page.images)

    full_text = "\n".join(pages_text)

    if column_detected:
        warnings.append("Multi-column layout detected — ATS may misorder sections")
    if table_detected:
        warnings.append("Tables detected — content may be lost or scrambled in ATS parse")
    if len(full_text.strip()) < 200:
        warnings.append("Very little text extracted — PDF may be image-based or heavily formatted")

    parse_integrity = _compute_parse_integrity(full_text, warnings)
    ats_preview = _strip_to_ats(full_text)

    audit = _build_audit(
        full_text,
        columns=column_detected,
        tables=table_detected,
        textboxes=None,          # not detectable in PDF
        images=image_count,
        pages=page_count,
    )

    return {
        "text": full_text,
        "ats_preview": ats_preview,
        "parse_integrity": parse_integrity,
        "warnings": warnings,
        "formatting_audit": audit,
    }


def _extract_docx(file_bytes: bytes) -> dict:
    from docx import Document

    warnings: list[str] = []
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)

    # Check for text boxes (common DOCX formatting trap)
    body_xml = doc.element.body.xml
    textboxes = "txbx" in body_xml or "textbox" in body_xml.lower()
    if textboxes:
        warnings.append("Text boxes detected — content inside text boxes is invisible to most ATS")

    if len(full_text.strip()) < 200:
        warnings.append("Very little text extracted — content may be in text boxes or images")

    parse_integrity = _compute_parse_integrity(full_text, warnings)
    ats_preview = _strip_to_ats(full_text)

    audit = _build_audit(
        full_text,
        columns=False,           # column detection is PDF-geometry based
        tables=len(doc.tables) > 0,
        textboxes=textboxes,
        images=body_xml.count("graphicData"),
        pages=None,              # DOCX has no fixed page count pre-render
    )

    return {
        "text": full_text,
        "ats_preview": ats_preview,
        "parse_integrity": parse_integrity,
        "warnings": warnings,
        "formatting_audit": audit,
    }


def _build_audit(
    text: str,
    *,
    columns: bool,
    tables: bool,
    textboxes: bool | None,
    images: int,
    pages: int | None,
) -> list[dict]:
    """
    Named formatting checklist (audit gap #2). Each check: pass | fail | warn.
    Only checks we can actually compute — no fabricated confidence.
    """
    stripped = text.strip()
    # Non-ASCII ratio — decorative glyphs/icons get mangled by naive parsers.
    non_ascii = len(re.findall(r"[^\x00-\x7F]", stripped))
    ratio = non_ascii / max(1, len(stripped))

    audit: list[dict] = [
        {"check": "Single-column layout", "status": "fail" if columns else "pass",
         "detail": "Multiple columns can be read out of order by ATS parsers." if columns
                   else "No multi-column layout detected."},
        {"check": "No tables", "status": "fail" if tables else "pass",
         "detail": "Table contents are often scrambled or dropped by ATS parsers." if tables
                   else "No tables detected."},
        {"check": "Machine-readable text", "status": "fail" if len(stripped) < 200 else "pass",
         "detail": "Almost no extractable text — likely image-based." if len(stripped) < 200
                   else "Text extracts cleanly."},
        {"check": "Standard characters", "status": "warn" if ratio > 0.03 else "pass",
         "detail": "Many special characters/icons — these often turn into garbage in ATS text."
                   if ratio > 0.03 else "No unusual character usage."},
        {"check": "No images/graphics", "status": "warn" if images > 0 else "pass",
         "detail": f"{images} image(s) found — ATS cannot read content inside graphics."
                   if images > 0 else "No embedded images."},
    ]
    if textboxes is not None:
        audit.append({"check": "No text boxes", "status": "fail" if textboxes else "pass",
                      "detail": "Text-box content is invisible to most ATS." if textboxes
                                else "No text boxes detected."})
    if pages is not None:
        audit.append({"check": "Reasonable length", "status": "warn" if pages > 2 else "pass",
                      "detail": f"{pages} pages — recruiters and some ATS truncate long résumés."
                                if pages > 2 else f"{pages} page(s)."})
    return audit


def _strip_to_ats(text: str) -> str:
    """Simulate what a naive ATS parser sees: collapsed whitespace, no formatting."""
    text = re.sub(r"[^\x20-\x7E\n]", " ", text)
    text = re.sub(r" {2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _compute_parse_integrity(text: str, warnings: list[str]) -> float:
    score = 1.0
    penalty_map = {
        "Multi-column": 0.20,
        "Tables detected": 0.15,
        "Text boxes": 0.25,
        "Very little text": 0.40,
    }
    for warning in warnings:
        for key, penalty in penalty_map.items():
            if key in warning:
                score -= penalty
    return max(0.0, round(score, 2))
